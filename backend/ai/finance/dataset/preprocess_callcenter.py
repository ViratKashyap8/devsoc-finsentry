"""
Preprocess call center dataset (DOCX transcripts) to JSONL training format.

Converts Hugging Face call center dataset with DOCX transcripts to the
FinanceTrainExample JSONL format for LoRA fine-tuning.

Usage:
    python -m ai.finance.dataset.preprocess_callcenter \
        --input data/audio \
        --output data/finance/callcenter/combined_train.jsonl
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from docx import Document

from .format import FinanceTrainExample, save_dataset_jsonl


# Map conversation types to intents
CONV_TYPE_TO_INTENT = {
    "customer support": "inquiry",
    "customer service": "inquiry",
    "support": "inquiry",
    "billing": "account_info",
    "account": "account_info",
    "sale": "general",
    "sales": "general",
    "finance": "account_info",
    "account manager": "account_info",
    "pharma": "general",
}

# Map sentiment keywords to emotion labels
SENTIMENT_TO_EMOTION = {
    "cooperative": "calm",
    "patient": "calm",
    "frustrated": "frustrated",
    "confusion": "anxious",
    "anxious": "anxious",
    "angry": "angry",
    "upset": "upset",
    "neutral": "neutral",
    "positive": "calm",
    "negative": "frustrated",
}


def extract_transcript_lines(doc: Document) -> list[str]:
    """Extract timestamped transcript lines from DOCX."""
    lines = []
    in_transcript = False
    timestamp_pattern = re.compile(r"^\d{2}:\d{2}:\d{2}")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Detect start of transcript section
        if "transcription" in text.lower() or "transcript" in text.lower():
            in_transcript = True
            continue
        
        # If we see a timestamp pattern, we're in transcript
        if timestamp_pattern.match(text):
            in_transcript = True
            # Remove timestamp prefix
            text = timestamp_pattern.sub("", text).strip()
            if text:
                lines.append(text)
        elif in_transcript and text and not any(
            kw in text.lower() for kw in ["summary", "key points", "next steps", "primary purpose"]
        ):
            lines.append(text)
    
    return lines


def extract_metadata(doc: Document) -> dict[str, Any]:
    """Extract metadata from DOCX (conversation type, sentiment, etc.)."""
    metadata = {
        "conversation_type": None,
        "sentiment": None,
        "key_points": [],
    }
    
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    for i, text in enumerate(paragraphs):
        text_lower = text.lower()
        
        if text_lower == "conversation type" and i + 1 < len(paragraphs):
            metadata["conversation_type"] = paragraphs[i + 1].lower()
        elif text_lower == "customer sentiment" and i + 1 < len(paragraphs):
            metadata["sentiment"] = paragraphs[i + 1].lower()
        elif text_lower == "key points" and i + 1 < len(paragraphs):
            metadata["key_points"].append(paragraphs[i + 1])
    
    return metadata


def infer_intent(conv_type: str | None) -> str:
    """Infer intent from conversation type."""
    if not conv_type:
        return "general"
    
    conv_lower = conv_type.lower()
    for key, intent in CONV_TYPE_TO_INTENT.items():
        if key in conv_lower:
            return intent
    return "general"


def infer_emotion(sentiment: str | None) -> str:
    """Infer emotion from sentiment description."""
    if not sentiment:
        return "neutral"
    
    sentiment_lower = sentiment.lower()
    for key, emotion in SENTIMENT_TO_EMOTION.items():
        if key in sentiment_lower:
            return emotion
    return "neutral"


def infer_risk(intent: str, emotion: str) -> str:
    """Infer risk level from intent and emotion."""
    high_risk_intents = {"fraud_report", "complaint"}
    medium_risk_intents = {"dispute"}
    high_stress_emotions = {"angry", "frustrated", "upset", "anxious"}
    
    if intent in high_risk_intents:
        return "high"
    if emotion in high_stress_emotions and intent in medium_risk_intents:
        return "medium"
    if emotion in high_stress_emotions:
        return "medium"
    if intent in medium_risk_intents:
        return "low"
    return "none"


def extract_entities_from_text(text: str) -> list[dict[str, Any]]:
    """Extract basic financial entities using regex."""
    entities = []
    
    # Money amounts
    for match in re.finditer(r"\$[\d,]+(?:\.\d{2})?|\d+\s*(?:dollars?|USD)", text, re.IGNORECASE):
        entities.append({
            "text": match.group(),
            "label": "AMOUNT",
            "start": match.start(),
            "end": match.end(),
        })
    
    # Account numbers (4+ digits)
    for match in re.finditer(r"(?:account|card)(?:\s+(?:number|#))?\s*[:\s]?\s*(\d{4,})", text, re.IGNORECASE):
        entities.append({
            "text": match.group(1),
            "label": "ACCOUNT",
            "start": match.start(1),
            "end": match.end(1),
        })
    
    # Dates
    for match in re.finditer(
        r"\b(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{2,4})",
        text, re.IGNORECASE
    ):
        entities.append({
            "text": match.group(),
            "label": "DATE",
            "start": match.start(),
            "end": match.end(),
        })
    
    return entities


def extract_obligations_from_text(text: str) -> list[dict[str, Any]]:
    """Extract obligation patterns from text."""
    obligations = []
    text_lower = text.lower()
    
    obligation_patterns = [
        (r"(?:i'?ll|we'?ll|will)\s+(?:send|email|call|get back)", "follow_up"),
        (r"(?:i'?ll|we'?ll|will)\s+pay", "payment_promise"),
        (r"refund.*(?:processed|sent|issued)", "refund"),
        (r"(?:waive|remove).*(?:fee|charge)", "fee_waiver"),
        (r"escalat(?:e|ing)", "escalation"),
    ]
    
    for pattern, obl_type in obligation_patterns:
        if re.search(pattern, text_lower):
            obligations.append({
                "text": text[:100],
                "type": obl_type,
                "amount": None,
                "due_date": None,
            })
            break  # One obligation per utterance
    
    return obligations


def process_docx_file(docx_path: Path) -> list[FinanceTrainExample]:
    """Process a single DOCX file into training examples."""
    try:
        doc = Document(str(docx_path))
    except Exception as e:
        print(f"Error reading {docx_path}: {e}")
        return []
    
    metadata = extract_metadata(doc)
    transcript_lines = extract_transcript_lines(doc)
    
    if not transcript_lines:
        print(f"No transcript found in {docx_path}")
        return []
    
    intent = infer_intent(metadata["conversation_type"])
    emotion = infer_emotion(metadata["sentiment"])
    risk = infer_risk(intent, emotion)
    
    examples = []
    
    # Create examples from transcript lines (combine short lines)
    current_text = ""
    for line in transcript_lines:
        current_text += " " + line if current_text else line
        
        # Create example when we have enough text
        if len(current_text) >= 50 or line == transcript_lines[-1]:
            current_text = current_text.strip()
            if current_text:
                entities = extract_entities_from_text(current_text)
                obligations = extract_obligations_from_text(current_text)
                
                examples.append(FinanceTrainExample(
                    text=current_text,
                    intent=intent,
                    entities=entities,
                    obligations=obligations,
                    emotion=emotion,
                    risk_level=risk,
                    id=f"{docx_path.stem}_{len(examples)}",
                ))
            current_text = ""
    
    return examples


def preprocess_callcenter_dataset(
    input_dir: str | Path,
    output_path: str | Path,
) -> int:
    """
    Process all DOCX files in the call center dataset.
    
    Args:
        input_dir: Path to the dataset root (e.g., data/audio)
        output_path: Output JSONL file path
    
    Returns:
        Number of examples created
    """
    input_dir = Path(input_dir)
    output_path = Path(output_path)
    
    # Find all DOCX files
    docx_files = list(input_dir.rglob("*.docx"))
    print(f"Found {len(docx_files)} DOCX files")
    
    all_examples = []
    for docx_path in docx_files:
        examples = process_docx_file(docx_path)
        all_examples.extend(examples)
        print(f"  {docx_path.name}: {len(examples)} examples")
    
    # Save to JSONL
    output_path.parent.mkdir(parents=True, exist_ok=True)
    save_dataset_jsonl(output_path, all_examples)
    
    print(f"\nTotal: {len(all_examples)} examples saved to {output_path}")
    return len(all_examples)


def main():
    parser = argparse.ArgumentParser(description="Preprocess call center DOCX to JSONL")
    parser.add_argument(
        "--input",
        required=True,
        help="Input directory with DOCX files (e.g., data/audio)",
    )
    parser.add_argument(
        "--output",
        required=True,
        help="Output JSONL file path",
    )
    args = parser.parse_args()
    
    preprocess_callcenter_dataset(args.input, args.output)


if __name__ == "__main__":
    main()
